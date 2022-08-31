import sys
#sys.path.append("C:\\Users\\etc your directory)

print(sys.path)

from docx import Document
from docx.shared import Inches

#Open a blank document
document = Document()

#Create a title
document.add_heading('Document Title', 0)

#Create sections to the word document
document = Document()
sections = document.sections
print(sections)

#Unlink header and footers and assign variable for headers and footers
section = document.sections[0]
header = section.header
header.is_linked_to_previous = False
paragraph = header.paragraphs[0]
footer = section.footer
footer.is_linked_to_previous = False
fparagraph = footer.paragraphs[0]

#Add header text
hparagraph = header.paragraphs[0]
hparagraph.text = "Left Text\tCenter Text\tRight Text"
hparagraph.style = document.styles["Header"]

#Add footer text
fparagraph = footer.paragraphs[0]
fparagraph.text = "Left footer Text\tCenter Text\tRight footer Text"
fparagraph.style = document.styles["Footer"]

#Add a paragraph of text with bold and italic text
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

#Add a heading of a particular level
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

#Add a bullet point list and numbered list and image
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'Add an image in below here:')
document.add_picture('image.png', width=Inches(1.25))

#Add a list of data for using in a table
records = (
    (1, '111', 'Name1'),
    (2, '222', 'Name2'),
    (3, '333', 'Name1, Name2, Name2, and name3')
)

#Add a table with particular style
table = document.add_table(rows=1, cols=3)
#table.style = 'Light Shading Accent 1'
table.style = 'Light Grid Accent 2' #'Medium Grid 3 Accent 5'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

#add page breaks into the document
document.add_page_break()

document.add_paragraph(
    'Another paragraph after a page break.')

document.add_page_break()

print(len(sections))

#Change a section's orientation and width and height
from docx.enum.section import WD_SECTION, WD_ORIENT
new_section = document.add_section(WD_SECTION.ODD_PAGE)
new_section.start_type = WD_SECTION.ODD_PAGE
new_section.orientation, new_section.page_width, new_section.page_height
new_width, new_height = new_section.page_height, new_section.page_width
new_section.orientation = WD_ORIENT.LANDSCAPE
new_section.page_width = new_width
new_section.page_height = new_height
new_section.orientation, new_section.page_width, new_section.page_height

document.add_paragraph(
    'A paragraph on a horizontal orientation.')

print(len(sections))

#Create the margins for the document
from docx.shared import Inches
section.left_margin, section.right_margin
section.top_margin, section.bottom_margin
section.gutter
section.header_distance, section.footer_distance
section.left_margin = Inches(1.5)
section.right_margin = Inches(1)
section.left_margin, section.right_margin

document.add_page_break()

#Change the orientation back to portrait
another_new_section = document.add_section(WD_SECTION.EVEN_PAGE)
another_new_section.start_type = WD_SECTION.EVEN_PAGE
another_new_section.orientation, another_new_section.page_width, another_new_section.page_height
new_width, new_height = another_new_section.page_height, another_new_section.page_width
another_new_section.orientation = WD_ORIENT.PORTRAIT
another_new_section.page_width = new_width
another_new_section.page_height = new_height
another_new_section.orientation, another_new_section.page_width, another_new_section.page_height

document.add_paragraph(
    'A paragraph after orientation change again.')

document.add_page_break()

yet_another_new_section = document.add_section(WD_SECTION.ODD_PAGE)

print(len(sections))

#Add in indented text
paragraph = document.add_paragraph('This is the first line with indent.\nThis is the secondline with extra indent.\nThis is the third line.')
indent_paragraph_format = paragraph.paragraph_format
from docx.shared import Pt

indent_paragraph_format.left_indent
indent_paragraph_format.right_indent
indent_paragraph_format.first_line_indent

indent_paragraph_format.left_indent = Pt(36)
indent_paragraph_format.left_indent.pt

indent_paragraph_format.right_indent = Inches(0.25)
indent_paragraph_format.right_indent.pt

indent_paragraph_format.first_line_indent = Pt(-18)
indent_paragraph_format.first_line_indent.pt

document.add_page_break()

document.add_paragraph(
    'A paragraph with no spacing before and after.')

#Add space before and after paragraph text
paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.space_before
paragraph_format.space_before = Pt(12)
paragraph_format.space_before.pt

paragraph_format.space_after
paragraph_format.space_after = Pt(18)
paragraph_format.space_after.pt

document.add_paragraph(
    'A paragraph with 12pt spacing before and 18pt after.')

document.add_paragraph(
    'A paragraph with 12pt spacing before and 18pt after.')

document.save('demo.docx')